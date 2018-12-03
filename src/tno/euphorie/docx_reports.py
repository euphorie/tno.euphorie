# coding=utf-8
from euphorie.client.docx.compiler import DocxCompiler
from docx.enum.text import WD_BREAK
from euphorie.client.docx.views import ActionPlanDocxView
from pkg_resources import resource_filename
from euphorie.client import MessageFactory as _


class RIEDocxCompiler(DocxCompiler):
    _template_filename = resource_filename(
        'tno.euphorie',
        'templates/rie.docx',
    )

    def set_session_title_row(self, data):
        ''' This fills the workspace activity run with some text
        '''
        # request = self.request


        heading = self.t(_(
            "plan_report_intro_header", default=u"Introduction"))
        intro = self.t(_(
            "plan_report_intro_1",
            default=u"By filling in the list of questions, you have "
                    u"completed a risk assessment. This assessment is used to "
                    u"draw up an action plan. The progress of this action "
                    u"plan must be discussed annually and a small report must "
                    u"be written on the progress. Certain subjects might have "
                    u"been completed and perhaps new subjects need to be "
                    u"added."))

        doc = self.template
        doc.paragraphs[0].text = heading
        doc.paragraphs[0].style = "Heading 1"
        p = doc.add_paragraph(intro)
        p.add_run().add_break(WD_BREAK.PAGE)

        # self.template.paragraphs[0].text = data['heading']
        # txt = self.t(_("toc_header", default=u"Contents"))
        # par_contents = self.template.paragraphs[1]
        # par_contents.text = txt
        # par_toc = self.template.paragraphs[2]
        # par_before_break = self.template.paragraphs[3]
        # for nodes, heading in zip(data["nodes"], data["section_headings"]):
        #     if not nodes:
        #         continue
        #     par_toc.insert_paragraph_before(heading, style="TOC Heading 1")
        # survey = self.request.survey

        # footer_txt = self.t(
        #     _("report_identification_revision",
        #         default=u"This document was based on the OiRA Tool '${title}' "
        #                 u"of revision date ${date}.",
        #         mapping={"title": survey.published[1],
        #                  "date": formatDate(request, survey.published[2])}))

        # par_contents.insert_paragraph_before(
        #     formatDate(request, date.today()), style="Comment")
        # par_before_break.insert_paragraph_before("")
        # par_before_break.insert_paragraph_before(footer_txt, 'Footer')



class RIEActionPlanDocxView(ActionPlanDocxView):

    _compiler = RIEDocxCompiler


    def compile(self, data):
        '''
        '''
        self.set_session_title_row(data)
        self.set_body(data)